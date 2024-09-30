import streamlit as st
import stripe
import os
from io import BytesIO
import docx
import requests
import jwt
import datetime
from docx.oxml.ns import qn
import logging
import streamlit.components.v1 as components

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set full-width layout and page title
st.set_page_config(layout="wide", page_title="Corrector de Documentos DOCX")

# Configurar la clave secreta de Stripe desde Streamlit Secrets
stripe.api_key = st.secrets["STRIPE_SECRET_KEY"]  # Acceder a Stripe secret desde Streamlit secrets

# ID del producto para Stripe (almacenado en Streamlit Secrets)
PRODUCT_ID = st.secrets["STRIPE_PRODUCT_ID"]  # Asegúrate de agregar STRIPE_PRODUCT_ID en tus secrets

# Obtener JWT_SECRET desde Streamlit secrets
JWT_SECRET = st.secrets["JWT_SECRET"]

# Webhook secret para Stripe (agregar STRIPE_WEBHOOK_SECRET en Streamlit Secrets)
STRIPE_WEBHOOK_SECRET = st.secrets.get("STRIPE_WEBHOOK_SECRET", "")

# Función para generar un JWT para el success_url
def generate_jwt_token():
    payload = {
        "paid": True,
        "exp": datetime.datetime.utcnow() + datetime.timedelta(minutes=30)  # Expiración del token (30 minutos)
    }
    token = jwt.encode(payload, JWT_SECRET, algorithm="HS256")
    return token

# Función para verificar el token JWT
def verify_jwt_token(token):
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        if payload.get("paid"):
            return True
        return False
    except jwt.ExpiredSignatureError:
        st.error("El token ha expirado.")
        return False
    except jwt.InvalidTokenError:
        st.error("Token inválido.")
        return False

# Función para obtener el precio del producto
def get_price_for_product(product_id):
    try:
        prices = stripe.Price.list(product=product_id, active=True)
        if prices and prices['data']:
            return prices['data'][0].id
        else:
            st.error("No se encontraron precios para el producto.")
            return None
    except Exception as e:
        st.error(f"Error al obtener el precio: {e}")
        logger.error(f"Error al obtener el precio: {e}")
        return None

# Función para crear una sesión de pago con Stripe
def create_checkout_session(price_id):
    try:
        token = generate_jwt_token()  # Generar el token JWT para el success URL
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{
                "price": price_id,
                "quantity": 1,
            }],
            mode="payment",
            success_url=f"{st.secrets['APP_URL']}/?token={token}",  # Usar el token JWT en el success URL
            cancel_url=f"{st.secrets['APP_URL']}/?cancel=true",
        )
        logger.info("Sesión de pago creada exitosamente.")
        return session
    except Exception as e:
        st.error(f"Error al crear la sesión de pago: {e}")
        logger.error(f"Error al crear la sesión de pago: {e}")
        return None

# Función para conectar con la API de LanguageTool y aplicar correcciones (procesamiento por lotes)
def correct_text_with_languagetool(text, language):
    languagetool_url = "https://api.languagetool.org/v2/check"
    language_codes = {
        "en": "en-US",
        "es": "es",
        "fr": "fr",
        "de": "de",
        "pt": "pt",
        "it": "it",  # Agregado Italiano
        # Agregar más idiomas si es necesario
    }

    params = {
        'text': text,
        'language': language_codes.get(language, "en-US"),
        'level': 'picky',  # Usar el nivel "picky" para aplicar más correcciones
        'enabledCategories': 'grammar,style,typos',  # Habilitar correcciones gramaticales, de estilo y tipográficas
        'enabledRules': 'WHITESPACE_RULE,EN_UNPAIRED_BRACKETS,UPPERCASE_SENTENCE_START,WORDINESS,REDUNDANCY,MISSING_COMMA,COMMA_PARENTHESIS_WHITESPACE,DASH_RULE,EN_QUOTES,AGREEMENT_SENT_START,SENTENCE_FRAGMENT,MULTIPLICATION_SIGN,PASSIVE_VOICE,EXTRA_WHITESPACE,COMMA_BEFORE_CONJUNCTION,HYPHENATION_RULES,ITS_IT_IS,DUPLICATE_WORD,NO_SPACE_BEFORE_PUNCTUATION',
        'disabledCategories': 'COLLOQUIALISMS'  # Deshabilitar lenguaje coloquial para precisión científica
    }

    try:
        response = requests.post(languagetool_url, data=params)
        if response.status_code == 200:
            matches = response.json().get('matches', [])
            corrected_text = apply_corrections(text, matches)
            return corrected_text
        else:
            st.error("Error en la respuesta de la API de LanguageTool.")
            logger.error(f"LanguageTool API Error: {response.status_code} - {response.text}")
            return text
    except requests.exceptions.RequestException as e:
        st.error(f"Error en la solicitud de la API: {e}")
        logger.error(f"Error en la solicitud de la API de LanguageTool: {e}")
        return text

# Función para aplicar correcciones al texto
def apply_corrections(text, matches):
    corrections = []
    for match in matches:
        if match['replacements']:
            replacement = match['replacements'][0]['value']
            start_pos = match['offset']
            end_pos = start_pos + match['length']
            corrections.append((start_pos, end_pos, replacement))

    corrected_text = text
    offset = 0
    for start_pos, end_pos, replacement in sorted(corrections, key=lambda x: x[0]):
        corrected_text = corrected_text[:start_pos + offset] + replacement + corrected_text[end_pos + offset:]
        offset += len(replacement) - (end_pos - start_pos)
    return corrected_text

# Función para manejar el procesamiento de un documento completo
def process_document(document, language):
    full_text = []
    footnote_indices = []
    for paragraph in document.paragraphs:
        if not paragraph_contains_footnote_reference(paragraph):
            full_text.append(paragraph.text)
        else:
            footnote_indices.append(len(full_text))
            full_text.append(paragraph.text)

    text_to_correct = "\n".join(full_text)
    corrected_text = correct_text_with_languagetool(text_to_correct, language)

    if corrected_text == text_to_correct:
        st.warning("No se realizaron correcciones.")
        return document

    corrected_paragraphs = corrected_text.split("\n")

    for i, paragraph in enumerate(document.paragraphs):
        if not paragraph_contains_footnote_reference(paragraph):
            paragraph.text = corrected_paragraphs[i]

    return document

# Función para verificar si un párrafo contiene una referencia de nota al pie
def paragraph_contains_footnote_reference(paragraph):
    for run in paragraph.runs:
        for child in run._r:
            if child.tag == qn('w:footnoteReference'):
                return True
    return False

# Función para renderizar el botón de pago con Stripe
def render_payment_button(session_url):
    components.html(
        f"""
        <script>
            function redirectToStripe() {{
                window.location.href = "{session_url}";
            }}
        </script>
        <button onclick="redirectToStripe()" style="background-color:#6772E5; color:white; padding: 10px 20px; border:none; border-radius:5px; cursor:pointer; font-size:16px;">
            Pagar con Stripe
        </button>
        """,
        height=60,
    )

# Función principal de la aplicación Streamlit
def main():
    # Barra lateral con instrucciones y configuración de pago con Stripe
    with st.sidebar:
        st.header("Instrucciones")
        st.markdown("""
        **Esta aplicación te permite:**
        
        - Subir un documento DOCX.
        - Aplicar correcciones ortográficas y gramaticales al texto, excluyendo citas y nombres propios.
        - Preservar las notas al pie.
        
        **Idiomas soportados:**
        
        La aplicación puede corregir documentos en **Inglés, Español, Francés, Italiano, Alemán y Portugués**.
        
        **Longitud del documento:**
        
        - La aplicación tiene virtualmente sin límites en la longitud del documento.
        - Sin embargo, cuanto más largo sea el documento, más tiempo tomará procesarlo.
        - Por ejemplo, un documento de 100 páginas puede tardar hasta 10 minutos en procesarse.
        
        **Control de cambios:**
        
        Para enviar el documento con control de cambios, sigue estos pasos:
        
        1. **Abre Word** y ve a la pestaña **Revisar**.
        2. Usa la opción **Comparar Documentos** y selecciona:
            - El documento original que subiste.
            - El documento corregido descargado desde esta aplicación.
        3. Word mostrará el documento con control de cambios, donde puedes:
            - Aceptar o rechazar las correcciones propuestas.
        4. Guarda el documento con control de cambios y envíalo para su revisión.
        
        **Precio:**
        
        - Se cobra una tarifa fija de 4 USD por documento, independientemente de su longitud.
        
        **Descargo de responsabilidad:**
        
        Es responsabilidad del editor verificar todos los cambios antes de enviar el documento a su destinatario final.
        
        **Autor:** Dr. Moris Polanco (mp @ ufm.edu)
        """)

    # Columna principal para el contenido interactivo
    col_main, _ = st.columns([3, 1])  # Ajuste para centrar el contenido principal

    with col_main:
        st.title("Corrección Ortográfica y Gramatical de Documentos DOCX con Preservación de Notas al Pie")

        # Obtener parámetros de la URL
        query_params = st.experimental_get_query_params()
        token = query_params.get("token", [None])[0]
        cancel = query_params.get("cancel", [False])[0]

        if cancel:
            st.error("Pago cancelado. Por favor, intenta nuevamente.")
            # Limpiar los parámetros de la URL
            st.experimental_set_query_params()
            return

        if token:
            if verify_jwt_token(token):
                st.success("¡Pago completado! Ahora puedes subir y procesar tu documento.")
                # Limpiar los parámetros de la URL
                st.experimental_set_query_params()
            else:
                st.error("Token inválido. Por favor, completa el pago nuevamente.")
                # Limpiar los parámetros de la URL
                st.experimental_set_query_params()
                return
        else:
            st.warning("Debes completar el pago antes de usar la aplicación.")
            price_id = get_price_for_product(PRODUCT_ID)
            if price_id:
                session = create_checkout_session(price_id)
                if session:
                    render_payment_button(session.url)
            return  # No permitir continuar sin pago

        # Permitir la carga de archivos y selección de idioma solo después del pago
        language = st.selectbox("Selecciona el idioma del documento", ["en", "es", "fr", "it", "de", "pt"])

        uploaded_file = st.file_uploader("Sube un archivo DOCX", type="docx")

        if uploaded_file is not None:
            try:
                # Intentar abrir el archivo para verificar que es un DOCX válido
                document = docx.Document(uploaded_file)
                st.success("Documento cargado exitosamente.")
            except Exception as e:
                st.error("El archivo subido no es un DOCX válido o está dañado.")
                logger.error(f"Error al abrir el documento: {e}")
            else:
                if st.button("Enviar"):
                    try:
                        progress_bar = st.progress(0)
                        total_steps = 3  # Carga, procesamiento, descarga
                        progress_bar.progress(1 / total_steps)

                        # Procesar el documento
                        corrected_document = process_document(document, language)
                        progress_bar.progress(2 / total_steps)

                        # Guardar el documento corregido en BytesIO
                        corrected_file = BytesIO()
                        corrected_document.save(corrected_file)
                        corrected_file.seek(0)

                        progress_bar.progress(3 / total_steps)
                        progress_bar.empty()

                        st.download_button(
                            label="Descargar documento corregido",
                            data=corrected_file,
                            file_name="documento_corregido.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                        st.success("¡Documento procesado y descargado exitosamente!")
                    except Exception as e:
                        st.error(f"Ocurrió un error al procesar el documento: {e}")
                        logger.error(f"Error al procesar el documento: {e}")

if __name__ == "__main__":
    main()
